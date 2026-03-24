import os
import sys
import re
import time
import pandas as pd
import win32com.client
from docx import Document

# ==================================================
# 🔍 BUSCAR ARCHIVO EXCEL
# ==================================================
archivos_excel = [f for f in os.listdir() if f.lower().endswith(".xlsx")]

if not archivos_excel:
    print("❌ No se encontró ningún archivo .xlsx en la carpeta.")
    sys.exit(1)

nombre_archivo = archivos_excel[0]
print(f"📄 Excel detectado: {nombre_archivo}")

# ==================================================
# CONFIGURACIÓN
# ==================================================
plantilla_word = "plantilla sin autorizaciones.docx"
carpeta_pdf = "PDF SIN AUTORIZACIONES"
carpeta_temp = "temp_word"

os.makedirs(carpeta_pdf, exist_ok=True)
os.makedirs(carpeta_temp, exist_ok=True)

print("📝 Plantilla Word detectada:", plantilla_word)
print("⏳ Procesando...\n")

# ==================================================
# LEER EXCEL
# ==================================================
df = pd.read_excel(nombre_archivo)
total_registros = len(df)

# ==================================================
# FUNCIÓN DE REEMPLAZO
# ==================================================
def reemplazar_campos(doc, datos):
    patrones = {
        re.compile(rf"\b{re.escape(str(campo))}\b", re.IGNORECASE): str(valor)
        for campo, valor in datos.items()
    }

    for p in doc.paragraphs:
        for patron, valor in patrones.items():
            if patron.search(p.text):
                p.text = patron.sub(valor, p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for patron, valor in patrones.items():
                    if patron.search(cell.text):
                        cell.text = patron.sub(valor, cell.text)

# ==================================================
# LIMPIAR NOMBRE DE ARCHIVO
# ==================================================
def limpiar_nombre(texto):
    return re.sub(r'[\\/:*?"<>|]', '', str(texto)).strip()

# ==================================================
# INICIALIZAR WORD
# ==================================================
word = win32com.client.DispatchEx("Word.Application")
word.Visible = False

try:
    for i, fila in df.iterrows():
        datos = fila.to_dict()

        print(f"➡ Generando {i+1}/{total_registros} → contrato: {datos.get('campod', '')}")

        doc = Document(plantilla_word)
        reemplazar_campos(doc, datos)

        nombre_pdf = (
            f"{limpiar_nombre(datos.get('campod', ''))}_"
            f"{limpiar_nombre(datos.get('campoc', ''))}_"
            f"INDICADORES DE CALIDAD_"
            f"{limpiar_nombre(datos.get('campob', ''))}.pdf"
        )

        ruta_word = os.path.abspath(os.path.join(carpeta_temp, f"temp_{i}.docx"))
        ruta_pdf = os.path.abspath(os.path.join(carpeta_pdf, nombre_pdf))

        # Guardar Word
        doc.save(ruta_word)
        time.sleep(1)

        # Convertir a PDF
        documento = word.Documents.Open(ruta_word)
        documento.SaveAs(ruta_pdf, FileFormat=17)
        documento.Close(False)

        os.remove(ruta_word)

finally:
    word.Quit()

print("\n✅ PROCESO COMPLETADO")
print(f"📂 PDFs generados en: {carpeta_pdf}")
