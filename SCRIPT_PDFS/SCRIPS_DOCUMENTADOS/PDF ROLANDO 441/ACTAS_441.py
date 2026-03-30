import os
import shutil
import pandas as pd
from docx import Document
import win32com.client as win32

# -------------------------------------------------------
#  🔧 CONFIGURACIÓN INICIAL
# -------------------------------------------------------

excel_name = "excel.xlsx"
word_template = "plantilla.docx"
output_folder = "ACTAS_441_GENERADAS"

# Eliminar carpeta si ya existe
if os.path.exists(output_folder):
    try:
        shutil.rmtree(output_folder)
    except PermissionError:
        print("⚠ No se pudo borrar la carpeta porque un archivo está abierto. Ciérralo e intenta de nuevo.")
        exit()

os.makedirs(output_folder)

print(f"📄 Excel detectado: {excel_name}")
print(f"📝 Plantilla Word detectada: {word_template}")
print("⏳ Procesando...\n")

# -------------------------------------------------------
#  🔧 CARGAR EXCEL
# -------------------------------------------------------

df = pd.read_excel(excel_name)

# -------------------------------------------------------
#  🔧 FUNCIÓN DE REEMPLAZO ROBUSTO
# -------------------------------------------------------

def replace_text_in_runs(paragraph, replacements):
    """Reemplaza texto incluso cuando Word rompe palabras en runs."""
    for key, value in replacements.items():
        if key in paragraph.text:
            full_text = paragraph.text.replace(key, str(value))

            for run in paragraph.runs:
                run.text = ""

            paragraph.runs[0].text = full_text


def replace_all_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        replace_text_in_runs(paragraph, replacements)


def replace_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_all_paragraphs(cell.paragraphs, replacements)

# -------------------------------------------------------
#  🔧 INICIO DEL PROCESO
# -------------------------------------------------------

total_registros = len(df)
i=0 

for idx, row in df.iterrows():

       
    i = i+1
    contrato = row["campob"]  # campo2 → campob
    print(f"➡ Generando {i}/{total_registros} → contrato: {contrato}")


    # Crear mapa {{columna}} → valor
    replacements = {}
    for col in df.columns:
        tag = f"{{{{{col}}}}}"  # ejemplo: campoa -> {{campoa}}
        replacements[tag] = row[col]

    # Cargar plantilla
    doc = Document(word_template)

    # Reemplazar texto
    replace_all_paragraphs(doc.paragraphs, replacements)
    replace_in_tables(doc.tables, replacements)

    # Guardar DOCX temporal
    temp_doc = os.path.join(output_folder, f"temp_{idx}.docx")
    doc.save(temp_doc)

    # -------------------------------------------------------
    #   CONVERSIÓN A PDF
    # -------------------------------------------------------

    pdf_name = f"{contrato} - INDICADORES GESTION .pdf"
    pdf_path = os.path.join(output_folder, pdf_name)

    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False

        docx_file = word.Documents.Open(os.path.abspath(temp_doc))
        docx_file.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        docx_file.Close(False)
        word.Quit()

    except Exception as e:
        print("❌ Error generando PDF:", e)
        continue

    os.remove(temp_doc)

print("\n✅ PROCESO COMPLETADO")
print(f"📂 PDFs generados en: {output_folder}")
